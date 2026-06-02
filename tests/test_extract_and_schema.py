from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from docx_redline_writer import validate_edit_plan
from extract_docx_structure import extract_docx_structure
from ooxml_utils import normalized_hash


SCHEMA = ROOT / "schemas" / "edit-plan.schema.json"


def _sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Public sample paragraph.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Public table cell."
    doc.save(path)


class ExtractAndSchemaTest(unittest.TestCase):
    def test_extract_docx_structure_reports_paragraphs_tables_and_hashes(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.docx"
            _sample_docx(path)
            structure = extract_docx_structure(path)
            self.assertEqual(
                structure["paragraphs"][0]["text"],
                "Public sample paragraph.",
            )
            self.assertEqual(
                structure["paragraphs"][0]["normalized_text_hash"],
                normalized_hash("Public sample paragraph."),
            )
            self.assertEqual(
                structure["tables"][0]["rows"][0]["cells"][0]["paragraphs"][0]["text"],
                "Public table cell.",
            )

    def test_edit_plan_schema_accepts_supported_action(self):
        schema = json.loads(SCHEMA.read_text(encoding="utf-8"))
        self.assertEqual(schema["title"], "DOCX WPS Redline Edit Plan")
        plan = {
            "document_id": "schema-sample",
            "generated_at": "2026-06-02T12:00:00+08:00",
            "actions": [
                {
                    "action_id": "A001",
                    "action_type": "replace_sentence",
                    "target": {
                        "container_type": "paragraph",
                        "paragraph_index": 0,
                        "table_path": None,
                        "target_text": "Public sample paragraph.",
                        "context_before": "",
                        "context_after": "",
                        "normalized_text_hash": normalized_hash("Public sample paragraph."),
                    },
                    "replacement_text": "Public replacement paragraph.",
                    "comment": "Reviewer note.",
                }
            ],
        }
        self.assertEqual(validate_edit_plan(plan), [])

    def test_edit_plan_schema_rejects_missing_replacement(self):
        json.loads(SCHEMA.read_text(encoding="utf-8"))
        plan = {
            "document_id": "schema-sample",
            "generated_at": "2026-06-02T12:00:00+08:00",
            "actions": [
                {
                    "action_id": "A001",
                    "action_type": "replace_sentence",
                    "target": {
                        "container_type": "paragraph",
                        "paragraph_index": 0,
                        "table_path": None,
                        "target_text": "Public sample paragraph.",
                        "context_before": "",
                        "context_after": "",
                        "normalized_text_hash": normalized_hash("Public sample paragraph."),
                    },
                    "comment": "Reviewer note.",
                }
            ],
        }
        self.assertIn("actions[0] missing replacement_text", validate_edit_plan(plan))


if __name__ == "__main__":
    unittest.main()
