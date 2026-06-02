from __future__ import annotations

import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from docx_redline_writer import run_redline
from ooxml_utils import (
    COMMENTS_CONTENT_TYPE,
    NS,
    WORD_REL_COMMENTS,
    ensure_comments_content_type,
    ensure_comments_relationship,
    make_comment,
    make_comment_reference_run,
    new_comments_root,
    normalized_hash,
    parse_xml,
    qn,
    serialize_xml,
)
from validate_docx_redline import validate_docx


def _write_public_sample(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Project brief introduction.")
    doc.add_paragraph("The draft must be delivered by Friday.")
    doc.add_paragraph("Remove this temporary note.")
    doc.add_paragraph("Keep this anchor sentence.")
    doc.add_paragraph("Section Alpha")
    doc.add_paragraph("First range paragraph.")
    doc.add_paragraph("Second range paragraph.")
    doc.add_paragraph("Repeated target. Repeated target.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table item requires review."
    doc.save(path)


def _paragraph_target(text: str, index: int) -> dict:
    return {
        "container_type": "paragraph",
        "paragraph_index": index,
        "table_path": None,
        "target_text": text,
        "context_before": "",
        "context_after": "",
        "normalized_text_hash": normalized_hash(text),
    }


def _table_target(text: str) -> dict:
    return {
        "container_type": "table_cell",
        "paragraph_index": None,
        "table_path": {
            "table_index": 0,
            "row_index": 0,
            "cell_index": 0,
            "paragraph_index_in_cell": 0,
        },
        "target_text": text,
        "context_before": "",
        "context_after": "",
        "normalized_text_hash": normalized_hash(text),
    }


def _range_target(text: str, start: int, end: int) -> dict:
    return {
        "container_type": "paragraph_range",
        "paragraph_index": None,
        "paragraph_range": {
            "start_paragraph_index": start,
            "end_paragraph_index": end,
        },
        "table_path": None,
        "target_text": text,
        "context_before": "",
        "context_after": "",
        "normalized_text_hash": normalized_hash(text),
    }


def _action(action_id: str, action_type: str, target: dict, **kwargs) -> dict:
    action = {
        "action_id": action_id,
        "action_type": action_type,
        "target": target,
        "comment": kwargs.pop("comment", f"Comment for {action_id}."),
    }
    if "replacement_text" in kwargs:
        action["replacement_text"] = kwargs["replacement_text"]
    return action


def _plan(actions: list[dict]) -> dict:
    return {
        "document_id": "public-sample",
        "generated_at": "2026-06-02T12:00:00+08:00",
        "actions": actions,
    }


def _write_plan(path: Path, plan: dict) -> None:
    path.write_text(json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")


def _read_output_parts(path: Path):
    with zipfile.ZipFile(path) as package:
        return {
            name: parse_xml(package.read(name))
            for name in [
                "word/document.xml",
                "word/settings.xml",
                "word/comments.xml",
                "word/_rels/document.xml.rels",
                "[Content_Types].xml",
            ]
        }


def _all_text(root, xpath: str) -> str:
    return "".join(node.text or "" for node in root.xpath(xpath, namespaces=NS))


def _add_existing_comment_and_revision(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path) as package:
        parts = {name: package.read(name) for name in package.namelist()}
        infos = {info.filename: info for info in package.infolist()}

    document_root = parse_xml(parts["word/document.xml"])
    comments_root = (
        parse_xml(parts["word/comments.xml"])
        if "word/comments.xml" in parts
        else new_comments_root()
    )
    rels_root = (
        parse_xml(parts["word/_rels/document.xml.rels"])
        if "word/_rels/document.xml.rels" in parts
        else etree.Element(f"{{{NS['rel']}}}Relationships", nsmap={None: NS["rel"]})
    )
    content_types_root = parse_xml(parts["[Content_Types].xml"])

    paragraph = document_root.xpath(".//w:body//w:p", namespaces=NS)[0]
    start = etree.Element(qn("w", "commentRangeStart"))
    start.set(qn("w", "id"), "20")
    end = etree.Element(qn("w", "commentRangeEnd"))
    end.set(qn("w", "id"), "20")
    paragraph.append(start)
    paragraph.append(end)
    paragraph.append(make_comment_reference_run(20))
    comments_root.append(
        make_comment(
            20,
            "Existing",
            datetime_from_iso("2026-06-02T01:00:00+00:00"),
            "Existing comment stays.",
        )
    )

    revision = etree.Element(qn("w", "ins"))
    revision.set(qn("w", "id"), "50")
    revision.set(qn("w", "author"), "Existing")
    revision.set(qn("w", "date"), "2026-06-02T01:00:00Z")
    run = etree.SubElement(revision, qn("w", "r"))
    text = etree.SubElement(run, qn("w", "t"))
    text.text = " Existing revision stays."
    paragraph.append(revision)

    ensure_comments_relationship(rels_root)
    ensure_comments_content_type(content_types_root)

    parts["word/document.xml"] = serialize_xml(document_root)
    parts["word/comments.xml"] = serialize_xml(comments_root)
    parts["word/_rels/document.xml.rels"] = serialize_xml(rels_root)
    parts["[Content_Types].xml"] = serialize_xml(content_types_root)

    with zipfile.ZipFile(docx_path, "w") as package:
        for name, data in parts.items():
            package.writestr(infos.get(name, name), data)


def datetime_from_iso(value: str):
    from datetime import datetime

    return datetime.fromisoformat(value.replace("Z", "+00:00"))


class RedlineWriterTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls._fixture_tmp = tempfile.TemporaryDirectory()
        cls.fixture_docx = Path(cls._fixture_tmp.name) / "public_sample.docx"
        _write_public_sample(cls.fixture_docx)

    @classmethod
    def tearDownClass(cls):
        cls._fixture_tmp.cleanup()

    def test_all_action_types_create_real_revisions_and_comments(self):
        original_bytes = self.fixture_docx.read_bytes()
        range_text = "First range paragraph.\nSecond range paragraph."
        actions = [
            _action(
                "A001",
                "replace_sentence",
                _paragraph_target("The draft must be delivered by Friday.", 1),
                replacement_text="The draft must be delivered by Monday.",
                comment="Deadline changed after review.",
            ),
            _action(
                "A002",
                "delete_sentence",
                _paragraph_target("Remove this temporary note.", 2),
                comment="Temporary note should be deleted.",
            ),
            _action(
                "A003",
                "insert_sentence_after",
                _paragraph_target("Keep this anchor sentence.", 3),
                replacement_text=" Added sentence for reviewers.",
                comment="Add follow-up context.",
            ),
            _action(
                "A004",
                "replace_clause",
                _range_target(range_text, 5, 6),
                replacement_text="Combined range replacement.",
                comment="Range replaced as one clause.",
            ),
            _action(
                "A005",
                "comment_only",
                _table_target("Table item requires review."),
                comment="Table comment only.",
            ),
        ]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            plan_path = tmp_dir / "edit_plan.json"
            output = tmp_dir / "public_sample_redlined.docx"
            log_path = tmp_dir / "redline_log.json"
            _write_plan(plan_path, _plan(actions))

            run_redline(
                input_docx=self.fixture_docx,
                edit_plan_path=plan_path,
                output_docx=output,
                log_path=log_path,
                random_seed=42,
                now="2026-06-02T12:00:00+08:00",
            )

            self.assertEqual(self.fixture_docx.read_bytes(), original_bytes)
            parts = _read_output_parts(output)
            document_root = parts["word/document.xml"]
            settings_root = parts["word/settings.xml"]
            comments_root = parts["word/comments.xml"]
            rels_root = parts["word/_rels/document.xml.rels"]
            content_types_root = parts["[Content_Types].xml"]

            self.assertEqual(len(document_root.xpath(".//w:del", namespaces=NS)), 4)
            self.assertEqual(len(document_root.xpath(".//w:ins", namespaces=NS)), 3)
            self.assertIn(
                "The draft must be delivered by Friday.",
                _all_text(document_root, ".//w:delText"),
            )
            self.assertIn(
                "The draft must be delivered by Monday.",
                _all_text(document_root, ".//w:ins//w:t"),
            )
            comment_texts = [
                _all_text(comment, ".//w:t")
                for comment in comments_root.xpath(".//w:comment", namespaces=NS)
            ]
            self.assertIn("Deadline changed after review.", comment_texts)
            self.assertTrue(
                settings_root.xpath("./w:revisionView/w:markup[text()='1']", namespaces=NS)
            )
            self.assertTrue(
                rels_root.xpath(
                    "./rel:Relationship[@Type=$type]",
                    namespaces=NS,
                    type=WORD_REL_COMMENTS,
                )
            )
            self.assertTrue(
                content_types_root.xpath(
                    "./ct:Override[@PartName='/word/comments.xml' and @ContentType=$content_type]",
                    namespaces=NS,
                    content_type=COMMENTS_CONTENT_TYPE,
                )
            )
            self.assertTrue(
                document_root.xpath(
                    ".//w:r[w:commentReference]/w:rPr/w:rStyle[@w:val='CommentReference']",
                    namespaces=NS,
                )
            )

            report = validate_docx(output, log_path, expected_author="Reviewer")
            self.assertTrue(report["valid"], report)
            log = json.loads(log_path.read_text(encoding="utf-8"))
            self.assertEqual(log["summary"]["applied"], 5)
            self.assertEqual(log["author"], "Reviewer")

    def test_author_can_be_overridden(self):
        actions = [
            _action(
                "A001",
                "replace_sentence",
                _paragraph_target("The draft must be delivered by Friday.", 1),
                replacement_text="The draft must be delivered by Monday.",
            )
        ]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            plan_path = tmp_dir / "edit_plan.json"
            output = tmp_dir / "out.docx"
            log_path = tmp_dir / "log.json"
            _write_plan(plan_path, _plan(actions))
            run_redline(
                input_docx=self.fixture_docx,
                edit_plan_path=plan_path,
                output_docx=output,
                log_path=log_path,
                author="QA Reviewer",
            )
            parts = _read_output_parts(output)
            deletions = parts["word/document.xml"].xpath(".//w:del", namespaces=NS)
            comments = parts["word/comments.xml"].xpath(".//w:comment", namespaces=NS)
            self.assertEqual(deletions[0].get(qn("w", "author")), "QA Reviewer")
            self.assertEqual(comments[0].get(qn("w", "author")), "QA Reviewer")

    def test_unresolved_target_does_not_guess(self):
        target = _paragraph_target("The draft must be delivered by Friday.", 1)
        target["target_text"] = "Missing target text."
        target["normalized_text_hash"] = normalized_hash("Missing target text.")
        actions = [
            _action(
                "A001",
                "replace_sentence",
                target,
                replacement_text="The draft must be delivered by Monday.",
            )
        ]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            plan_path = tmp_dir / "bad_plan.json"
            output = tmp_dir / "out.docx"
            log_path = tmp_dir / "log.json"
            _write_plan(plan_path, _plan(actions))
            run_redline(
                input_docx=self.fixture_docx,
                edit_plan_path=plan_path,
                output_docx=output,
                log_path=log_path,
            )
            parts = _read_output_parts(output)
            self.assertFalse(parts["word/document.xml"].xpath(".//w:del", namespaces=NS))
            log = json.loads(log_path.read_text(encoding="utf-8"))
            self.assertEqual(log["summary"]["unresolved"], 1)

    def test_hash_mismatch_is_unresolved(self):
        target = _paragraph_target("The draft must be delivered by Friday.", 1)
        target["normalized_text_hash"] = normalized_hash("Different text.")
        actions = [
            _action(
                "A001",
                "replace_sentence",
                target,
                replacement_text="The draft must be delivered by Monday.",
            )
        ]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            plan_path = tmp_dir / "bad_hash.json"
            output = tmp_dir / "out.docx"
            log_path = tmp_dir / "log.json"
            _write_plan(plan_path, _plan(actions))
            run_redline(
                input_docx=self.fixture_docx,
                edit_plan_path=plan_path,
                output_docx=output,
                log_path=log_path,
            )
            log = json.loads(log_path.read_text(encoding="utf-8"))
            self.assertEqual(log["actions"][0]["reason"], "target_text_hash_mismatch")

    def test_repeated_target_is_unresolved(self):
        actions = [
            _action(
                "A001",
                "replace_sentence",
                _paragraph_target("Repeated target.", 7),
                replacement_text="Unique replacement.",
            )
        ]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            plan_path = tmp_dir / "repeated.json"
            output = tmp_dir / "out.docx"
            log_path = tmp_dir / "log.json"
            _write_plan(plan_path, _plan(actions))
            run_redline(
                input_docx=self.fixture_docx,
                edit_plan_path=plan_path,
                output_docx=output,
                log_path=log_path,
            )
            log = json.loads(log_path.read_text(encoding="utf-8"))
            self.assertEqual(log["actions"][0]["reason"], "target_text_not_unique")

    def test_paragraph_range_cross_parent_boundary_is_unresolved(self):
        text = "Second range paragraph.\nTable item requires review."
        actions = [
            _action(
                "A001",
                "replace_clause",
                _range_target(text, 6, 8),
                replacement_text="Should not apply.",
            )
        ]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            plan_path = tmp_dir / "cross_parent.json"
            output = tmp_dir / "out.docx"
            log_path = tmp_dir / "log.json"
            _write_plan(plan_path, _plan(actions))
            run_redline(
                input_docx=self.fixture_docx,
                edit_plan_path=plan_path,
                output_docx=output,
                log_path=log_path,
            )
            log = json.loads(log_path.read_text(encoding="utf-8"))
            self.assertEqual(
                log["actions"][0]["reason"],
                "paragraph_range_crosses_parent_boundary",
            )

    def test_output_must_not_overwrite_source(self):
        actions = [
            _action(
                "A001",
                "comment_only",
                _paragraph_target("Project brief introduction.", 0),
            )
        ]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            sample = tmp_dir / "sample.docx"
            _write_public_sample(sample)
            plan_path = tmp_dir / "plan.json"
            log_path = tmp_dir / "log.json"
            _write_plan(plan_path, _plan(actions))
            with self.assertRaises(ValueError):
                run_redline(
                    input_docx=sample,
                    edit_plan_path=plan_path,
                    output_docx=sample,
                    log_path=log_path,
                )

    def test_existing_comments_and_revisions_are_preserved(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            sample = tmp_dir / "sample.docx"
            _write_public_sample(sample)
            _add_existing_comment_and_revision(sample)
            actions = [
                _action(
                    "A001",
                    "comment_only",
                    _paragraph_target("Project brief introduction. Existing revision stays.", 0),
                    comment="New comment.",
                )
            ]
            plan_path = tmp_dir / "plan.json"
            output = tmp_dir / "out.docx"
            log_path = tmp_dir / "log.json"
            _write_plan(plan_path, _plan(actions))
            run_redline(
                input_docx=sample,
                edit_plan_path=plan_path,
                output_docx=output,
                log_path=log_path,
            )
            parts = _read_output_parts(output)
            comments = parts["word/comments.xml"].xpath(".//w:comment", namespaces=NS)
            ids = {int(comment.get(qn("w", "id"))) for comment in comments}
            self.assertIn(20, ids)
            self.assertTrue(any(comment_id > 20 for comment_id in ids))
            revisions = parts["word/document.xml"].xpath(".//w:ins", namespaces=NS)
            revision_ids = {int(revision.get(qn("w", "id"))) for revision in revisions}
            self.assertIn(50, revision_ids)


if __name__ == "__main__":
    unittest.main()
